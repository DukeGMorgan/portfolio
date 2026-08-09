"""Build an ATS-safe DOCX resume for Duke Morgan.

ATS rules honoured here: no tables, no text boxes, no columns, no images,
no header/footer content, standard section headings, plain bullet lists,
and a common font. Everything is linear body text so a parser reads it
top-to-bottom in the right order.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0F, 0x3D, 0x56)   # deep slate blue
BODY = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x44, 0x4A, 0x50)

FONT = "Calibri"


def set_margins(doc, inches=0.45):
    for s in doc.sections:
        s.top_margin = Inches(inches)
        s.bottom_margin = Inches(inches)
        s.left_margin = Inches(inches + 0.05)
        s.right_margin = Inches(inches + 0.05)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.font.color.rgb = BODY
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0


def para(doc, text="", size=10, bold=False, italic=False, color=BODY,
         align=None, space_before=0, space_after=3, caps=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        r.font.name = FONT
        if caps:
            r.font.all_caps = True
    return p


def rich(doc, parts, size=10, space_before=0, space_after=3, align=None):
    """parts = [(text, bold, italic, color), ...]"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        r.font.name = FONT
    return p


def bottom_rule(p, color="0F3D56", size=6):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    bdr.append(bottom)
    pPr.append(bdr)


def section(doc, title):
    p = para(doc, title, size=10.5, bold=True, color=ACCENT,
             space_before=6, space_after=3, caps=True)
    for r in p.runs:
        r.font.spacing = Pt(1)
    bottom_rule(p)
    return p


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT
    r.font.color.rgb = BODY
    return p


def lead_bullet(doc, lead, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(lead)
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.name = FONT
    r1.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.name = FONT
    r2.font.color.rgb = BODY
    return p


def role(doc, title, org, meta):
    rich(doc, [(title, True, False, BODY)], size=10.5,
         space_before=6, space_after=0)
    rich(doc, [(org, True, False, ACCENT), ("  ·  ", False, False, MUTED),
               (meta, False, False, MUTED)], size=9.5, space_after=3)


def build():
    doc = Document()
    set_margins(doc)
    style_base(doc)

    # ---- Header -------------------------------------------------------
    p = para(doc, "DUKE MORGAN, RN, BSHA", size=21, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    for r in p.runs:
        r.font.spacing = Pt(1.2)

    para(doc, "Director of Clinical Systems  |  Clinical & Operational AI, "
              "Enterprise EMR, Automation at Scale",
         size=10.5, bold=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    p = para(doc,
             "Nashville, TN Region  ·  931-384-0389  ·  Dukeisrn@me.com  ·  "
             "linkedin.com/in/duke-morgan-rn-bsha-029b50b",
             size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    bottom_rule(p, size=8)

    # ---- Summary ------------------------------------------------------
    section(doc, "Summary")
    para(doc,
         "Registered nurse turned healthcare technology leader, with 20+ years spanning "
         "emergency department floors, enterprise EMR implementation, and applied AI "
         "architecture. I own clinical systems strategy for a multi-state fertility care "
         "network: multiple enterprise EMR platforms, hundreds of clinical users, and a "
         "multimillion-dollar annual vendor portfolio.")
    para(doc,
         "I also build. I put AI to work implementing and troubleshooting enterprise EMR "
         "platforms and automating the day-to-day work around them, and I write the SQL, "
         "Python, and automation that runs on it. Very few people in healthcare IT hold an "
         "active clinical license and a production codebase at the same time. That combination "
         "is why my roadmaps survive contact with real clinics.")

    # ---- AI leadership (front-loaded, deliberately above competencies) --
    section(doc, "AI Strategy & Platform Leadership")
    for lead, body in [
        ("Governance first. ",
         "Authored the two-lane data policy that decides what an AI system is allowed to see: "
         "PHI-bearing work stays on BAA-covered or fully local inference, while schema, code, "
         "and architecture work routes to commercial models. The boundary is enforced at the "
         "routing layer, so compliance never depends on someone remembering the rule under "
         "deadline."),
        ("AI strategy as an operating rhythm. ",
         "Run a standing weekly review to identify which workflows should be automated: "
         "clinical work inside the EMR, the administrative work around it, and the routine "
         "daily tasks that consume staff time. The measure is not a flagship system; it is "
         "whether analysts, schedulers, and clinicians each spend less of the week on work "
         "that did not need a person."),
        ("Multi-agent orchestration in production. ",
         "Architected and operate a fleet of specialized AI agents behind a model routing "
         "gateway to Azure AI Foundry, with local on-premise inference reserved for "
         "PHI-sensitive workloads. Agents are scoped by role and data lane rather than given "
         "blanket access."),
        ("AI applied to clinical operations, not demos. ",
         "Shipped AI into the work the organization actually does: schema intelligence over a "
         "600+ table clinical database, legacy EMR migration agents, and multi-site "
         "administrative automation."),
        ("Institutional memory that compounds. ",
         "Built a shared, persistent knowledge base every agent reads from and writes to, so "
         "operational lessons and root-cause findings accumulate instead of being rediscovered."),
    ]:
        lead_bullet(doc, lead, body, size=9.5)

    # ---- Competencies -------------------------------------------------
    section(doc, "Core Competencies")
    comps = [
        ("Leadership & Strategy:  ",
         "Enterprise EMR Product Ownership · Clinical Systems Roadmap · Multi-Site Deployment · "
         "Vendor & Contract Negotiation · Executive & C-Suite Stakeholder Management · Budget Ownership"),
        ("Clinical Informatics:  ",
         "EMR/EHR Implementation & Optimization · Clinical Workflow Design · Clinical Decision "
         "Support · CPOE · HL7 / FHIR Interoperability · Legacy EMR Data Migration · Go-Live Leadership"),
        ("AI, Automation & Engineering:  ",
         "Multi-Agent AI Orchestration · Azure AI Foundry · Model Routing Gateways · Local / "
         "On-Premise LLM Deployment for PHI Isolation · Retrieval-Augmented Generation · "
         "AI-Assisted Workflow Automation · Automation Opportunity Discovery"),
        ("Compliance & Security:  ",
         "HIPAA & PHI Governance · BAA-Governed AI Architecture · HITRUST-Aligned Design · "
         "Entra ID / SSO · Role-Based Access Control · Audit & Access Remediation · "
         "Joint Commission / CMS Readiness"),
        ("Technical:  ",
         "MS SQL Server (600+ table clinical schemas) · Python · PowerShell · .NET · "
         "React / TypeScript · REST APIs · Git · Azure · Supabase / PostgreSQL · "
         "Windows Enterprise Administration · Microsoft 365"),
    ]
    for lead, body in comps:
        rich(doc, [(lead, True, False, ACCENT), (body, False, False, BODY)],
             size=9.5, space_after=3)

    # ---- Experience ---------------------------------------------------
    section(doc, "Professional Experience")

    role(doc, "Director of Clinical Systems",
         "Inception Fertility Ventures, LLC", "Nashville, TN  ·  Jun 2017 – Present")
    para(doc,
         "Clinical systems strategy and product ownership for the largest fertility care network "
         "in North America: multi-state operations serving hundreds of clinical users.",
         size=9.5, italic=True, color=MUTED, space_after=4)

    for b in [
        "Own the full product lifecycle for multiple enterprise EMR platforms: roadmap, release "
        "governance, vendor alignment, and optimization across the workflows clinicians use to "
        "deliver patient care.",
        "Manage a multimillion-dollar annual vendor contract portfolio. Negotiations, SLA "
        "enforcement, and renewal strategy, consistently expanding capability without "
        "expanding spend.",
        "Direct cross-functional delivery across clinical operations, engineering, external "
        "vendors, and the executive team. Sustained 95% on-time delivery through platform "
        "upgrades, cloud migrations, and compliance programs.",
        "Built the business case for enterprise AI in a PHI environment and established a "
        "BAA-governed Azure AI Foundry pathway for AI-assisted clinical operations work, with "
        "the data-governance policy settled before any tooling was adopted.",
        "Charter and own the clinical systems roadmap: multiple server migrations and platform "
        "conversions, system-wide EMR application updates, and the introduction of new systems "
        "integrated with the EMR to improve clinical efficiency. Each is architected so AI "
        "tooling touches schema and metadata only, never patient data.",
        "Serve as the escalation point for defects the vendor cannot reproduce, including a "
        "null-dereference that made affected patient charts both un-openable and invisible to "
        "search.",
    ]:
        bullet(doc, b, size=9.5)

    # ---- Prior roles --------------------------------------------------
    role(doc, "Clinical Implementation Specialist, Perioperative Systems",
         "MEDHOST", "Nov 2016 – Jun 2017")
    bullet(doc, "Led surgical workflow optimization for multi-hospital EMR implementations, "
                "including gap analysis and future-state design across OR, pre-op, and PACU "
                "environments.", size=9.5)
    bullet(doc, "Configured perioperative modules and delivered clinician training that "
                "measurably improved go-live adoption.", size=9.5)

    role(doc, "Clinical Application Analyst, Corporate Deployment",
         "Community Health Systems", "Mar 2016 – Oct 2016")
    bullet(doc, "Developed enterprise EMR workflow templates to corporate clinical and "
                "regulatory standards, and supported go-lives onsite across emergency, "
                "surgical, and med/surg departments.", size=9.5)

    role(doc, "Support & Implementation Clinical Analyst, RN",
         "MEDHOST", "Jan 2013 – Mar 2016")
    bullet(doc, "Conducted clinical workflow analysis across ED, OR, and Med/Surg, configured "
                "EMR systems to clinical best practice, and partnered with bedside clinicians "
                "to validate build decisions before go-live.", size=9.5)

    # ---- Selected build work (own section, not buried in Experience) ---
    section(doc, "Selected Architecture & Build Work")

    for lead, body in [
        ("Clinical Intelligence Orchestration Platform. ",
         "Architected an enterprise multi-agent AI system routed through a model gateway to "
         "Azure AI Foundry under an active BAA. A two-lane topology keeps PHI-bearing work on "
         "BAA-covered or fully local inference while routing schema and code work to commercial "
         "models, establishing a HIPAA-aligned pathway for AI in clinical operations."),
        ("Unified Clinical Record Access Platform. ",
         "Delivered a .NET 8 and React/TypeScript application on Azure consolidating four legacy "
         "EMR platforms into one read-only clinical viewer behind Entra ID SSO. Migrated 3M+ "
         "clinical records, 770K+ documents, and 80K+ treatment cycles, eliminated "
         "shared-credential access across all sites, and closed a cross-clinic PHI exposure with "
         "a global authorization filter."),
        ("Multi-Agent Workflow Automation Platform. ",
         "Built the platform the team's operational automation runs on: purpose-built agents "
         "that absorb repetitive non-clinical work, sharing a persistent knowledge base so "
         "context carries across sessions. Each agent reviews its own sessions to turn "
         "recurring manual work into reusable skills, built for one agent and then rolled out "
         "to the rest."),
        ("Cross-Model Agent Collaboration Layer. ",
         "A shared working room where agents built on different model families work the same "
         "problem together, proposing, challenging and handing off to one another rather than "
         "running in isolation. Model choice becomes a per-task decision, and no single vendor "
         "is a single point of failure."),
        ("EMR Administration & Data Operations Console. ",
         "Python/Flask multi-site administration tool spanning every site database in the "
         "estate: chart merge, "
         "bulk provisioning, drug library management, and permission governance, with a "
         "plan-then-apply safety model that re-validates against live state and refuses to run "
         "on drift."),
        ("Clinical Schema Intelligence Assistant. ",
         "AI-assisted query assistant over a 600+ table clinical SQL Server schema, scoped to "
         "schema and metadata only. Cut root-cause analysis time and shortened analyst onboarding."),
    ]:
        lead_bullet(doc, lead, body, size=9.5)

    # ---- Selected findings (own section) ------------------------------
    section(doc, "Selected Findings & Remediations")

    for b in [
        "Audited application permissions across two production sites and executed a governed "
        "revocation of an over-provisioned module, verified by a rollback proving zero "
        "unintended change. Discovered in the process that the vendor's procedure writes no "
        "audit record for partial permission changes. The gap affects every such change made "
        "through the vendor UI, not just automated ones.",
        "Traced a data-integrity fault in a laboratory results migration to a non-unique record "
        "key silently collapsing duplicate groups that held differing payloads. Re-keyed the "
        "pipeline and recovered the records that would otherwise have been lost.",
        "Identified and closed a publicly readable patient document storage bucket, replacing it "
        "with authenticated, row-level-secured access.",
    ]:
        bullet(doc, b, size=9.5)

    # ---- Earlier practice ---------------------------------------------
    section(doc, "Earlier Clinical Practice")

    rich(doc, [("Emergency Department Charge Nurse", True, False, BODY),
               ("  ·  Williamson Medical Center  ·  Apr 2010 – Jan 2013",
                False, False, MUTED)], size=9.5, space_before=3, space_after=1)
    para(doc, "Led ED nursing operations as CPOE superuser, streamlining documentation "
              "workflows and routing operational findings back to IT.",
         size=9.5, space_after=3)

    rich(doc, [("Emergency Department Charge Nurse & Staff Nurse", True, False, BODY),
               ("  ·  Southern Hills Medical Center (HCA)  ·  Jul 2005 – Jan 2010",
                False, False, MUTED)], size=9.5, space_before=3, space_after=1)
    para(doc, "Managed emergency nursing operations and supported EMR template optimization "
              "for Joint Commission and regulatory compliance.", size=9.5, space_after=3)

    # ---- Education ----------------------------------------------------
    section(doc, "Education")
    rich(doc, [("Bachelor of Science, Healthcare Administration (BSHA)", True, False, BODY),
               ("  ·  Western Governors University, 2025", False, False, MUTED)], size=9.5)
    rich(doc, [("Associate of Applied Science, Nursing (ASN)", True, False, BODY),
               ("  ·  Excelsior College, 2007", False, False, MUTED)], size=9.5)

    # ---- Licensure ----------------------------------------------------
    section(doc, "Licensure & Certifications")
    bullet(doc, "Registered Nurse (RN), Tennessee / Multistate Compact License, Active", size=9.5)
    bullet(doc, "Microsoft SQL Server, Specialization Certificate, Microsoft via Coursera, 2026",
           size=9.5)
    bullet(doc, "AI Agent Fundamentals with Azure AI Foundry, Course Certificate, "
                "Microsoft via Coursera, 2026", size=9.5)

    out = (r"C:\Users\DukeMorgan\.dmorgan_Portifolio_creation\site\resume"
           r"\Duke-Morgan-Resume.docx")
    doc.save(out)
    print("WROTE", out)


if __name__ == "__main__":
    build()
